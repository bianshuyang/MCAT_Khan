from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions
from selenium import webdriver
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException
import urllib.request
from docx.shared import RGBColor
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from urllib.request import *
from docx import Document
from docx.shared import Inches
from urllib import request
import os
import time
import traceback
def mimickuser(xpath):
        e = driver.find_element_by_xpath(xpath)
        location = e.location
        driver.execute_script("window.scrollTo("+str(location.get('x'))+', '+str(location.get('y'))+')')   #("window.scrollTo(0, Y)") 
def clickitem(item,description):
        str_error = False
        while str_error == False:
            try:
                item.click()
                str_error = True
                print(description + 'has been selected!')
                break
            except:
                str_error = False
                time.sleep(5)
                print('sleeping 5 seconds')
def passageitem(item,description):
        str_error = False
        while str_error == False:
            try:
                ans = item.text
                str_error = True
                print(description + 'has been selected!')
                break
                print (ans)
            except:
                str_error = False
                time.sleep(3)
                print('sleeping 3 seconds')
def get_hint():
        firsthint = '/html/body/div[4]/div[3]/div/div[2]/div/div/div[3]/div/div[3]/div[2]/div[1]/div/div[3]/div/div/div[1]/div/div/div[3]/div/div[1]/div[1]/button'
        WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, firsthint)))
        mimickuser(firsthint)
        firsthint = driver.find_element_by_xpath(firsthint)
        clickitem(firsthint,'the first hint ')
        flag = True
        while flag == True:
                try:
                        resthint = '/html/body/div[4]/div[3]/div/div[2]/div/div/div[3]/div/div[3]/div[2]/div[1]/div/div[3]/div/div/div[1]/div/div/div[1]/div/form/div/div[2]/div/button'
                        WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, resthint)))
                        mimickuser(resthint)
                        resthint = driver.find_element_by_xpath(resthint)
                        clickitem(resthint,'the rest hint ')
                except:
                        flag = False
                        break
def get_article_to_docx(title):
    from docx import Document
    num = 1
    document = Document()
    text = '1'
    txts = []
    while (text == '')== False:
        xp = ('/html/body/div[4]/div[3]/div/div[2]/div/div/div[3]/div/div[3]/div[2]/div[1]/div/div[3]/div/div/div[1]/div/div/div[1]/div/form/div/div[1]/div/div['+str(num)+']')
        try:
            text = (driver.find_element_by_xpath(xp))
            print (text)
            if '.png' in text.get_attribute('innerHTML'):
                pure = text.text
                txts.append(pure)
                document.add_paragraph(pure)
                ps = text.get_attribute('innerHTML')
                img_link = ps[ps.index('https://cdn.kastatic.org'):ps.index('.png')+4]
                request.urlretrieve(img_link,filename=pure[0:9]+'.png')
                document.add_picture(pure[0:9]+'.png',width=Inches(6.3))
            else:
                pure = text.text
                txts.append(pure)
                document.add_paragraph(pure)
                document.save(title+'.docx')
                print(title+'docx'+'saved!')
        except:
            document.save(title+'.docx')
            print(title+'.docx saved!')
            text = ''
            return(num)
            break           
        num = num+1
def get_hint_to_docx(title):
        from docx import Document
        num = 1
        document = Document()
        text = '1'
        while (text == '')== False:
            xp = ('/html/body/div[4]/div[3]/div/div[2]/div/div/div[3]/div/div[3]/div[2]/div[1]/div/div[3]/div/div/div[1]/div/div/div[1]/div/form/div/div[2]/div/div['+str(num)+']')
            try:
                text = (driver.find_element_by_xpath(xp))
                pure = text.text
                document.add_paragraph(pure)
            except Exception as e:
                print(e)
                print(num)
                print('end of hint')
                text = ''
                break
            num = num+1
        document.save(title+'_hint.docx')
        print (num)
def getting_choices_answers(passage_title):### This will output 'ans.txt' as the answer file. But this won't be processed until the next function was run.
        content = []
        grand_auto = 1
        while grand_auto<=30:
                try:
                        get_hint()
                        return_value = get_article_to_docx(passage_title+str(grand_auto))
                        get_article_to_docx(passage_title+str(grand_auto))
                        get_hint_to_docx(passage_title+str(grand_auto))
                        choice_number = 1
                        while choice_number<=4:
                                choice = ('/html/body/div[4]/div[3]/div/div[2]/div/div/div[3]/div/div[3]/div[2]/div[1]/div/div[3]/div/div/div[1]/div/div/div[1]/div/form/div/div[1]/div/div['+str(return_value-1)+']/div/div/div/fieldset/ul/li['+str(choice_number)+']')
                        ####choice = ('/html/body/div[4]/div[3]/div/div[2]/div/div/div[3]/div/div[3]/div[2]/div[1]/div/div[3]/div/div/div[1]/div/div/div[1]/div/form/div/div[1]/div/div[10]/div/div/div/fieldset/ul/li[1]')
                                WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, choice)))
                                mimickuser(choice)
                                choice = driver.find_element_by_xpath(choice)
                                clickitem(choice,'one possible answer choice')
                                check = '/html/body/div[4]/div[3]/div/div[2]/div/div/div[3]/div/div[3]/div[2]/div[1]/div/div[3]/div/div/div[2]/div[1]/div[2]/button'
                                WebDriverWait(driver, 30).until(expected_conditions.visibility_of_element_located((By.XPATH, check)))
                                check = driver.find_element_by_xpath(check)
                                clickitem(check,'an option to check an answer')
                                correctness = ('/html/body/div[4]/div[3]/div/div[2]/div/div/div[3]/div/div[3]/div[2]/div[1]/div/div[3]/div/div/div[1]/div/div/div[1]/div/form/div/div[1]/div/div['+str(return_value-1)+']')
                                WebDriverWait(driver, 30).until(expected_conditions.visibility_of_element_located((By.XPATH, correctness)))
                                correctness= driver.find_element_by_xpath(correctness)
                                (passageitem(correctness,'answer choice description'))
                                content.append(correctness.text)
                                try:
                                        nextpage = ('/html/body/div[4]/div[3]/div/div[2]/div/div/div[3]/div/div[3]/div[2]/div[1]/div/div[3]/div/div/div[2]/div[1]/div[2]/button/div')
                                        ####pag  = ('/html/body/div[4]/div[3]/div/div[2]/div/div/div[3]/div/div[3]/div[2]/div[1]/div/div[3]/div/div/div[2]/div[1]/div[2]/button')
                                        WebDriverWait(driver, 2).until(expected_conditions.visibility_of_element_located((By.XPATH, nextpage)))
                                        nextpage = driver.find_element_by_xpath(nextpage)
                                        clickitem(nextpage,'going to nextpage')
                                        windows = driver.window_handles
                                        driver.switch_to.window(windows[0])
                                        time.sleep(5)
                                        break
                                except Exception as e:
                                        print(e)
                                        tryagain = '/html/body/div[4]/div[3]/div/div[2]/div/div/div[3]/div/div[3]/div[2]/div[1]/div/div[3]/div/div/div[2]/div[1]/div[2]/button'
                                        ####ain = '/html/body/div[4]/div[3]/div/div[2]/div/div/div[3]/div/div[3]/div[2]/div[1]/div/div[3]/div/div/div[2]/div[1]/div[2]/button'
                                        ####ain = '/html/body/div[4]/div[3]/div/div[2]/div/div/div[3]/div/div[3]/div[2]/div[1]/div/div[3]/div/div/div[2]/div[1]/div[2]/button'
                                        WebDriverWait(driver, 30).until(expected_conditions.visibility_of_element_located((By.XPATH, tryagain)))
                                        tryagain = driver.find_element_by_xpath(tryagain)
                                        clickitem(tryagain,'going to nextpage')
                                        choice_number = choice_number+1
                        grand_auto = grand_auto +1
                except Exception as e:
                        print(e)
                        print('maybe you have finishsed all articles!')
                        break
        with open('ans.txt','w') as f:
                f.write('\n'.join(content))
def answer_mod(title):
        with open('ans.txt','r') as f:
                a = f.readlines()
        final = []
        for i in a:
                i= i.replace('\n','')
                i =i.replace('(','')
                i= i.replace(')','')
                correct = ['Choice A, Checked, Correct','Choice B, Checked, Correct','Choice C, Checked, Correct','Choice D, Checked, Correct']
                for k in correct:
                        if i==k:
                                i = i.replace('Choice ','')
                                i = i.replace(', Checked, Correct','')
                                final.append(i)
        with open(title+'_ans_mod.txt','w') as f:
                f.write('\n'.join(final))
import os
b = os.path.abspath(os.path.join(os.getcwd(), "../.."))
c = os.getcwd()
current = c.replace(b,'')
current = current.replace('\\','/')
current = current.replace('/exercises','')
with open('exercises.txt','r') as f:
    a = f.readlines()
links = []
for i in a:
    if current in i:
        i = i.replace('\n','')
        links.append(i)
links = list(set(links))
for link in links:
        driver = webdriver.Chrome('C://Users//1//AppData//Local//Google//Chrome//Application//chromedriver.exe')
        passage_title = link[link.index('/e/')+3:]
        driver.get(link);
        getting_choices_answers(passage_title)
        answer_mod(passage_title)
        driver.close()
