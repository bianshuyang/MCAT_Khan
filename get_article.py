import urllib.request
import os
from docx import Document
from docx.shared import Inches
def get_grand_link(link):
    response = urllib.request.urlopen(link)
    print('getting response',link)
    a = response.read().decode('utf-8')
    print ('reading response')
    with open('1.txt','w',encoding = 'utf-8',errors='ignore') as f:
        f.write(a)
        print ('writing response')
    with open('1.txt','r',encoding = 'utf-8') as f:
        a = f.readlines()
    a = ''.join(a)
    renew = 'perseusContent'
    a = a[a.index(renew)+len(renew):]
    with open('1.txt','w',encoding = 'utf-8',errors = 'ignore') as f:
        f.write(a)
        print('get grand link success')
def imgdoclink(text):
    with open(text,'r',encoding= 'utf-8',errors = 'ignore') as f:
        a = f.readlines()
    a = ''.join(a)
    imgstart = 'widgets\\'
    imgend = 'thumbnailUrls'
    b = a[a.index(imgstart)+len(imgstart):a.index(imgend)]
    with open('img.txt','w',encoding= 'utf-8',errors = 'ignore') as f:
        f.write(b)
    c = a[0:a.index(imgstart)]
    with open('this_article.txt','w',encoding= 'utf-8',errors = 'ignore') as f:
        f.write(c)

def imglink():
    links = []
    with open('img.txt','r',encoding = 'utf-8',errors = 'ignore') as f:
        a = f.readlines()
    a = ''.join(a)
    while 'https://cdn.kastatic.org' in a:
        start = 'https://cdn.kastatic.org'
        if '.png' in a and '.svg' in a:
            a1 = a.index(".png")
            a2 = a.index('.svg')
            if a1<a2:
                end = '.png'
                lk = a[a.index(start):a.index(end)+4]
                links.append(lk)
                a = a.replace(lk,'')
            else:
                end = '.svg'
                lk = a[a.index(start):a.index(end)+4]
                links.append(lk)
                a = a.replace(lk,'')
        else:
            if '.png' in a:
                end = '.png'
                lk = a[a.index(start):a.index(end)+4]
                links.append(lk)
                a = a.replace(lk,'')
            else:
                end = '.svg'
                lk = a[a.index(start):a.index(end)+4]
                links.append(lk)
                a = a.replace(lk,'')
    return(links)
def downloadsvg():
    links = imglink()
    from urllib import request
    for i in links:
        if '.svg' in i:
            request.urlretrieve(i,filename=str(links.index(i))+'.svg')
        else:
            request.urlretrieve(i,filename=str(links.index(i))+'.png')
        print('convert success',i)
    import os
    from svglib.svglib import svg2rlg
    from reportlab.graphics import renderPDF, renderPM
    a = os.listdir(os.getcwd())
    for i in a:
        if '.svg' in i:
            try:
                drawing = svg2rlg(i)
                renderPM.drawToFile(drawing, i[0]+'.png', fmt="PNG")
                print ('convert picture right')
                os.remove(i)
            except:
                import shutil
                
def removeall(title):
    import os
    loc = os.getcwd()
    things = os.listdir(os.getcwd())
    things.remove('get_article.py')
    things.remove('articles.txt')
    for i in things:
        if '.txt' in i and (i=='articles.txt')==False:
            os.remove(i)
        if '.png' in i:
            os.remove(i)
    print ('cleaning your folder complete!')    

with open('articles.txt','r') as f:
    links = f.readlines()
for i in links:
    a = os.getcwd()
    a = a.replace('E:\\MCAT_Khan\\','')
    a = a.replace('articles','')
    a = a.replace('\\','/')
    q = i.replace('/a/','/')
    print(a,q)
    if a in q:
        title = i[i.index('/a/')+3:]
        title = title.replace("\n",'')
        get_grand_link(i)
        imgdoclink('1.txt')
        imglink()
        downloadsvg()
        key = 0
        from docx import Document
        from docx.shared import Inches
        with open('this_article.txt','r',encoding = 'utf-8') as f:
            a = f.readlines()
        a = ''.join(a)
        a = a.replace('\\\\n\\\\n','\n')
        a = a.replace('\\\\n','\n')
        start = '":\"'
        try:
            a = a.replace(a[:a.index(start)],'')
        except:
            a = a
        with open('new_article.txt','w',encoding = 'utf-8') as f:
            f.write(a)
        with open('new_article.txt','r',encoding = 'utf-8') as f:
            a = f.readlines()
        document = Document()
        def getxieti(i):
            flag = 0
            start = i.index('*')
            i = i[0:start]+'彝'+i[start+1:]
            start = i.index('彝')
            end = i.index('*')
            q = i[start:end+1]
            q = q.replace('彝','')
            q = q.replace('*','')
            if flag ==0:
                p = document.add_paragraph(i[0:start])
            else:
                p.add_run(i[0:start])
            p.add_run(q).italic =True
            remain = i[end+1:]
            if '*' not in remain:
                p.add_run(remain)
            else:
                getxieti(remain)
                flag =1
        def getxiahua(i):
            start = i.index('__')
            i = i[0:start]+'彝'+i[start+1:]
            start = i.index('彝')
            end = i.index('*')
            q = i[start:end+1]
            q = q.replace('彝','')
            q = q.replace('*','')
            p = document.add_paragraph(i[0:start])
            p.add_run(q).underline = True
            remain = i[end+1:]
            if '__' not in remain:
                p.add_run(remain)
            else:
                getxieti(remain)
        for i in a:
            i = i.replace('":"[{\"content\":\"','')
            if '#' in i:
                i = i.replace('#','')
                document.add_heading(i, level=1)
            else:
                if '*' in i:
                    try:
                        getxieti(i)
                    except:
                        document.add_paragraph(i)
                else:
                    if '☃' in i:
                        try:
                            document.add_picture(str(key)+'.png',width=Inches(6.3))
                            key = key+1
                        except:
                            print ('local_pictures_used_up!')
                    else:
                        if '__' in i:
                            try:
                                getxiahua(i)
                            except:
                                document.add_paragraph(i)
                        else:
                            document.add_paragraph(i)
        document.save(title+'.docx')
        print (title+'document saved!')
        import time
        removeall(title)
